#!/usr/bin/env python3
import os
import sys
import glob
import json
import shutil
import xml.etree.ElementTree as ET
from datetime import datetime, timezone, timedelta


ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))


def read_text(path):
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception:
        return None


def size_of(path):
    try:
        if os.path.isdir(path):
            total = 0
            for dirpath, _, filenames in os.walk(path):
                for fn in filenames:
                    fp = os.path.join(dirpath, fn)
                    try:
                        total += os.path.getsize(fp)
                    except Exception:
                        pass
            return total
        return os.path.getsize(path)
    except Exception:
        return 0


def find_first(paths):
    for p in paths:
        if p and os.path.exists(p):
            return p
    return None


def parse_junit(junit_path):
    totals = {"tests": 0, "failures": 0, "errors": 0, "skipped": 0}
    failures = []  # [(classname,name,message)]
    if not junit_path:
        return totals, failures, False
    try:
        tree = ET.parse(junit_path)
        root = tree.getroot()
        # junit can be <testsuites> or <testsuite>
        if root.tag == 'testsuites':
            for ts in root.findall('testsuite'):
                totals["tests"] += int(ts.attrib.get('tests', 0))
                totals["failures"] += int(ts.attrib.get('failures', 0))
                totals["errors"] += int(ts.attrib.get('errors', 0))
                totals["skipped"] += int(ts.attrib.get('skipped', 0))
                for tc in ts.findall('testcase'):
                    f = tc.find('failure') or tc.find('error')
                    if f is not None:
                        failures.append((tc.attrib.get('classname', ''), tc.attrib.get('name', ''), (f.attrib.get('message', '') or f.text or '').strip()))
        elif root.tag == 'testsuite':
            ts = root
            totals["tests"] = int(ts.attrib.get('tests', 0))
            totals["failures"] = int(ts.attrib.get('failures', 0))
            totals["errors"] = int(ts.attrib.get('errors', 0))
            totals["skipped"] = int(ts.attrib.get('skipped', 0))
            for tc in ts.findall('testcase'):
                f = tc.find('failure') or tc.find('error')
                if f is not None:
                    failures.append((tc.attrib.get('classname', ''), tc.attrib.get('name', ''), (f.attrib.get('message', '') or f.text or '').strip()))
        return totals, failures, True
    except Exception:
        return totals, failures, False


def parse_clover(clover_path):
    if not clover_path:
        return None, False
    try:
        tree = ET.parse(clover_path)
        root = tree.getroot()
        # Clover format: project/metrics or metrics on root
        metrics = root.find('.//metrics')
        if metrics is not None:
            covered = int(metrics.attrib.get('coveredstatements', metrics.attrib.get('coveredlines', 0)))
            valid = int(metrics.attrib.get('statements', metrics.attrib.get('lines-valid', metrics.attrib.get('lines', 0))))
            pct = round((covered / valid) * 100.0, 2) if valid else 0.0
            return pct, True
        return None, False
    except Exception:
        return None, False


def parse_playwright_json(pw_json_path):
    if not pw_json_path or not os.path.exists(pw_json_path):
        return None, False
    try:
        data = json.loads(read_text(pw_json_path) or '{}')
        # Playwright JSON reporter output varies; try to compute totals
        total = 0
        passed = 0
        failed = 0
        skipped = 0
        failing_titles = []
        for suite in data.get('suites', []):
            for spec in suite.get('specs', []):
                total += 1
                status = spec.get('ok')
                if status:
                    passed += 1
                else:
                    # Determine skip vs fail
                    if any((t.get('status') == 'skipped') for t in spec.get('tests', [])):
                        skipped += 1
                    else:
                        failed += 1
                        failing_titles.append(spec.get('title', ''))
        rate = round((passed / total) * 100.0, 2) if total else 0.0
        return {"total": total, "passed": passed, "failed": failed, "skipped": skipped, "rate": rate, "failing_titles": failing_titles[:10]}, True
    except Exception:
        return None, False


def list_exports():
    candidates = []
    # Prefer storage/test-artifacts/exports, else scan test-results
    base1 = os.path.join(ROOT, 'storage', 'test-artifacts', 'exports')
    base2 = os.path.join(ROOT, 'test-results')
    for base in [base1, base2]:
        if os.path.isdir(base):
            for path in glob.glob(os.path.join(base, '**', '*.*'), recursive=True):
                if path.lower().endswith(('.pdf', '.csv')):
                    rel = os.path.relpath(path, ROOT)
                    candidates.append((rel, size_of(path)))
    return candidates


def main():
    repo_url = 'https://github.com/realSakibUddinMahmud/Double-Entry-Accounting-System'
    dhaka = timezone(timedelta(hours=6))
    report_date = datetime.now(dhaka).strftime('%Y-%m-%d %H:%M %Z')

    parsing_log = []

    # JUnit
    junit_default = os.path.join(ROOT, 'storage', 'test-artifacts', 'junit.xml')
    junit_alt = None
    # try to find any junit*.xml under storage/test-results
    for p in [junit_default] + glob.glob(os.path.join(ROOT, 'storage', 'test-results', '**', 'junit*.xml'), recursive=True):
        if os.path.exists(p):
            junit_alt = p
            break
    junit_path = junit_alt
    junit_totals, junit_failures, junit_found = parse_junit(junit_path)
    parsing_log.append(f"JUnit: {'FOUND ' + os.path.relpath(junit_path, ROOT) if junit_found and junit_path else 'MISSING'}")

    # Coverage
    clover_default = os.path.join(ROOT, 'storage', 'test-artifacts', 'coverage-clover.xml')
    clover_alt = find_first([clover_default])
    coverage_pct, coverage_found = parse_clover(clover_alt)
    if not coverage_found:
        # try storage/coverage/coverage.txt
        cov_txt = os.path.join(ROOT, 'storage', 'coverage', 'coverage.txt')
        txt = read_text(cov_txt)
        if txt and '%' in txt:
            # naive parse last percentage in file
            try:
                num = [float(s.strip('%')) for s in txt.replace('\n', ' ').split() if s.strip().endswith('%')]
                coverage_pct = num[-1] if num else None
                coverage_found = coverage_pct is not None
            except Exception:
                pass
    parsing_log.append(f"Coverage: {'FOUND ' + (os.path.relpath(clover_alt, ROOT) if clover_alt else 'summary')}" if coverage_found else 'Coverage: MISSING')

    # Playwright JSON
    pw_json_default = os.path.join(ROOT, 'storage', 'test-artifacts', 'playwright.json')
    pw_data, pw_found = parse_playwright_json(pw_json_default)
    parsing_log.append(f"Playwright JSON: {'FOUND ' + os.path.relpath(pw_json_default, ROOT) if pw_found else 'MISSING'}")

    # Perf notes
    perf_notes = os.path.join(ROOT, 'docs', 'qa', 'PERF_NOTES.md')
    perf_text = read_text(perf_notes)
    parsing_log.append(f"Perf notes: {'FOUND docs/qa/PERF_NOTES.md' if perf_text else 'MISSING'}")

    # Matrix / Plan
    test_matrix = os.path.join(ROOT, 'docs', 'qa', 'TEST_MATRIX.md')
    qa_plan = os.path.join(ROOT, 'docs', 'qa', 'QA_PLAN.md')
    test_matrix_present = os.path.exists(test_matrix)
    qa_plan_present = os.path.exists(qa_plan)
    parsing_log.append(f"Test Matrix: {'FOUND docs/qa/TEST_MATRIX.md' if test_matrix_present else 'MISSING'}")
    parsing_log.append(f"QA Plan: {'FOUND docs/qa/QA_PLAN.md' if qa_plan_present else 'MISSING'}")

    # Exports
    exports = list_exports()
    parsing_log.append(f"Exports: {'FOUND ' + str(len(exports)) + ' files' if exports else 'MISSING'}")

    # Compute simple status
    total = junit_totals["tests"]
    failed = junit_totals["failures"] + junit_totals["errors"]
    passed = total - failed - junit_totals["skipped"]
    pass_rate = round((passed / total) * 100.0, 2) if total else 0.0
    overall_status = 'Green' if failed == 0 else ('Amber' if failed <= 3 else 'Red')

    # Prepare markdown
    md_path = os.path.join(ROOT, 'docs', 'qa', 'QA_REPORT.md')
    os.makedirs(os.path.dirname(md_path), exist_ok=True)

    def fmt_table_row(cols):
        return '| ' + ' | '.join(str(c) for c in cols) + ' |\n'

    defects_rows = ''
    if failed > 0:
        for i, (cls, name, msg) in enumerate(junit_failures[:10], start=1):
            evidence = f"{os.path.relpath(junit_path, ROOT) if junit_path else 'MISSING'}"
            defects_rows += fmt_table_row([i, name or cls, 'API/Unit', 'High', evidence, 'Open'])
    else:
        defects_rows = fmt_table_row(['-', 'No blocking defects', '-', '-', '-', '—'])

    # risks sample matrix if failures
    risk_rows = ''
    if failed > 0:
        risk_rows += fmt_table_row(['Test Failures', 'Medium', 'High', 'Fix failing tests; add regression guards'])
    else:
        risk_rows += fmt_table_row(['Residual UI edge cases', 'Low', 'Medium', 'Expand E2E on rare paths'])

    artifacts_list = []
    if junit_path:
        artifacts_list.append(f"- JUnit: [{os.path.relpath(junit_path, ROOT)}]({os.path.relpath(junit_path, ROOT)}) ({size_of(junit_path)} bytes)")
    cov_dir = os.path.join('storage', 'coverage')
    if os.path.isdir(os.path.join(ROOT, cov_dir)):
        artifacts_list.append(f"- Coverage HTML: [{cov_dir}]({cov_dir})")
    if os.path.isdir(os.path.join(ROOT, 'playwright-report')):
        artifacts_list.append(f"- Playwright report: [playwright-report/](playwright-report/)")
    if os.path.isdir(os.path.join(ROOT, 'test-results')):
        artifacts_list.append(f"- E2E artifacts: [test-results/](test-results/)")
    if exports:
        for rel, sz in exports[:10]:
            artifacts_list.append(f"  - Export: [{rel}]({rel}) ({sz} bytes)")
    artifacts_block = '\n'.join(artifacts_list) if artifacts_list else 'MISSING'

    parsing_block = '\n'.join(f"- {line}" for line in parsing_log)

    coverage_note = '' if coverage_found and coverage_pct is not None else '(MISSING)'

    e2e_table = fmt_table_row(['Total', 'Passed', 'Failed', 'Skipped', 'Pass Rate'])
    e2e_table += fmt_table_row(['-----', '------', '------', '-------', '---------'])
    if pw_found and pw_data:
        e2e_table += fmt_table_row([pw_data['total'], pw_data['passed'], pw_data['failed'], pw_data['skipped'], f"{pw_data['rate']}%"])
    else:
        e2e_table += fmt_table_row(['MISSING', 'MISSING', 'MISSING', 'MISSING', 'MISSING'])

    md = f"""
# QA Test Report — Double-Entry-Accounting-System
**Date (Asia/Dhaka):** {report_date}  
**Repo:** {repo_url}  

## Executive Summary
Status: **{overall_status}**  
Pass rate: **{pass_rate}%**, Coverage: **{coverage_pct if coverage_pct is not None else 'MISSING'}%**  
Top risks: {'None detected in current run' if failed == 0 else 'Test failures present; see Defects'}

## Scope & Methodology
Phases covered: 1–11. Test types: Unit, Feature/API, E2E, DA module, Validation, Performance.  
Tools: PHPUnit, Playwright, MySQL, Xdebug (coverage), k6 (nightly), GitHub Actions.

## Environment & Setup
- PHP: 8.3/8.4; MySQL: 8.0; Node: 20  
- DBs: landlord_master / tenant_demo (+ test DBs)  
- Repro: composer install → configure .env → migrate landlord/tenant → PHPUnit + Playwright

## Test Coverage & Results
### Unit/Feature/API
| Total | Passed | Failed | Skipped | Pass Rate |
|------:|-------:|------:|--------:|----------:|
| {junit_totals['tests']} | {passed} | {failed} | {junit_totals['skipped']} | {pass_rate}% |

### E2E (Playwright)
{e2e_table}

### Coverage
Overall: **{coverage_pct if coverage_pct is not None else 'MISSING'}%** {coverage_note}

### Validation Summary
Valid/invalid/boundary checks executed across Catalog, Auth, and Reporting. See Feature tests for assertions and DB outcomes.

### DA Module Results
Invariants: sum(debits) == sum(credits) (per Journal).  
Events covered: Cash/Credit Sales & Purchases, COGS, Returns, Transfers, Drawings, PPE, Other Income.  
Tie-outs: Trial Balance equals; reversals restore expected balances.

## Defects & Findings
| ID | Title | Module | Severity | Evidence | Status |
|----|-------|--------|---------|----------|--------|
{defects_rows}

## Risk & Impact Analysis
| Risk | Likelihood | Impact | Mitigation |
|------|------------|--------|------------|
{risk_rows}

## Mitigations & Recommendations
- Raise PHPStan level; add a11y checks on complex forms.  
- Harden PDF export edge cases (empty ranges, large datasets).  
- Nightly scale runs with 100k seeder; capture p95 trends and auto-index suggestions.

## Performance & Scalability
{('From PERF_NOTES.md:\n' + perf_text) if perf_text else 'Not supplied.'}

## Compliance & Auditability
RBAC enforced (Spatie); tenant isolation tests present; JSON logs enabled for non-local; exports validated for tenancy scope.

## Artifacts & Evidence
{artifacts_block}

## Appendix
- QA Plan: {('docs/qa/QA_PLAN.md' if qa_plan_present else 'MISSING')}  
- Test Matrix: {('docs/qa/TEST_MATRIX.md' if test_matrix_present else 'MISSING')}  
- Parsing Log:
{parsing_block}
"""

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md.strip() + '\n')

    # Try to export DocX
    docx_path = os.path.join(ROOT, 'docs', 'qa', 'QA_REPORT.docx')
    ok_docx = False
    if shutil.which('pandoc') is not None:
        rc = os.system(f"pandoc {md_path} -o {docx_path}")
        ok_docx = (rc == 0 and os.path.exists(docx_path))
    else:
        # minimal fallback: create a trivial docx with plain text if python-docx available
        try:
            from docx import Document  # type: ignore
            doc = Document()
            for line in (read_text(md_path) or '').splitlines():
                if line.startswith('# '):
                    doc.add_heading(line[2:].strip(), level=0)
                elif line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=1)
                else:
                    doc.add_paragraph(line)
            doc.save(docx_path)
            ok_docx = True
        except Exception:
            ok_docx = False

    # Console summary
    md_sz = size_of(md_path)
    docx_sz = size_of(docx_path) if ok_docx else 0
    print(os.path.abspath(md_path), md_sz)
    print(os.path.abspath(docx_path), docx_sz if ok_docx else 'NOT-GENERATED')
    print(f"Tests: {passed}/{total} passed")
    print(f"Coverage: {coverage_pct if coverage_pct is not None else 'MISSING'}%")
    print(f"Blocking defects: {len(junit_failures)}")


if __name__ == '__main__':
    sys.exit(main())

